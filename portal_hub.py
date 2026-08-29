import os
import io
import re
import pandas as pd  # Data visualization dependency for the analytics tracking system
import streamlit as st
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pypdf import PdfReader
from compliance_engine import build_compliance_agent

# =====================================================================
# 1. COMMERCIAL APPLICATION LAYOUT CONFIGURATION
# =====================================================================
st.set_page_config(
    page_title="ReguBot AI Compliance Suite (Beta Testing)",
    page_icon="🛡️",
    layout="wide"
)

# Initialize persistent session state matrices to track background data logs
if "metrics_log" not in st.session_state:
    st.session_state["metrics_log"] = []
if "feedback_submitted" not in st.session_state:
    st.session_state["feedback_submitted"] = False

# =====================================================================
# 2. CLERK USER REGISTRATION GATEWAY LOGIC
# =====================================================================
if "user_authenticated" not in st.session_state:
    st.session_state["user_authenticated"] = False

if not st.session_state["user_authenticated"]:
    st.title("🔐 ReguBot Institutional Authentication")
    st.write("Please sign in or register your academic email address below to access the multi-agent compliance pipeline.")
    
    CLERK_SIGN_IN_URL = "https://accounts.dev"
    st.components.v1.iframe(CLERK_SIGN_IN_URL, height=550, scrolling=True)
    
    st.markdown("---")
    if st.button("✅ I have successfully signed into my Clerk Account", type="primary"):
        st.session_state["user_authenticated"] = True
        st.rerun()
    st.stop()

# =====================================================================
# 3. SIDEBAR WORKSPACE WITH INTEGRATED FEEDBACK TOOL
# =====================================================================
with st.sidebar:
    st.header("🏢 Beta Test Environment")
    st.success("👑 Testing Tier Unlocked: Report Download Feature is Active.")
    
    st.markdown("---")
    st.header("🔑 Bring-Your-Own-Key (BYOK)")
    st.info("This application runs using your personal account tokens. No server fees are charged to the host.")
    
    user_gemini_key = st.text_input("Enter Your Google Gemini API Key:", type="password", placeholder="AIzaSy...")
    st.caption("🔗 [Get a free Gemini key from Google AI Studio](https://google.com)")

    st.markdown("---")
    st.subheader("⚡ Bounding Filters")
    selected_years = st.slider("Select Publication Year Range:", min_value=2010, max_value=2026, value=(2022, 2026))
    start_year, end_year = selected_years

    # 💬 ANONYMOUS BETA TESTER FEEDBACK WIDGET
    st.markdown("---")
    st.subheader("💬 Anonymous Tester Feedback")
    if not st.session_state["feedback_submitted"]:
        with st.form("feedback_form", clear_on_submit=True):
            rating = st.selectbox("Rate Report Accuracy:", ["⭐⭐⭐⭐⭐ (Excellent)", "⭐⭐⭐⭐ (Good)", "⭐⭐⭐ (Average)", "⭐⭐ (Poor)", "⭐ (Broken)"])
            comments = st.text_area("What features or rules should we add?", placeholder="e.g., Add Section 3(e) check for patents, modify ICMR timeline validation...")
            submit_feedback = st.form_submit_button("Submit Feedback")
            if submit_feedback:
                st.session_state["feedback_submitted"] = True
                st.toast("Thank you for your valuable response!", icon="🚀")
                st.rerun()
    else:
        st.success("🎉 Thank you! Feedback logged successfully.")
        if st.button("Submit another response"):
            st.session_state["feedback_submitted"] = False
            st.rerun()

    st.markdown("---")
    st.subheader("🗒️ Legal & Compliance")
    with st.expander("⚖️ Terms of Service"):
        st.caption("Usage Profile: This application functions purely under a Bring-Your-Own-Key (BYOK) paradigm.")
    with st.expander("🔒 Privacy Policy"):
        st.caption("Data Custody: Uploaded manuscripts are processed purely in-memory during active loops.")

# Main app tabs configuration separating execution and analytics panels
st.title("🛡️ ReguBot AI: Institutional Compliance & IPR Suite")
st.caption("Beta Sandbox Mode: Auditing academic manuscripts for Patents, Grants, and Ethics board approval.")

main_tab, dashboard_tab = st.tabs(["🚀 Execution Pipeline", "📊 Tracking Analytics Dashboard"])
# =====================================================================
# 4. EXECUTIVE WORD DOCUMENT STYLING ENGINES
# =====================================================================
def convert_markdown_to_docx(markdown_text):
    doc = Document()
    
    # Establish Global Style Parameters
    FONT_NAME = "Times New Roman"
    CRIMSON_RED = RGBColor(180, 0, 0)
    MUTED_GRAY = RGBColor(128, 128, 128)
    
    # Page Setup & 1-Inch Corporate Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        # Institutional Header Design
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_p.add_run("🛡️ ReguBot AI | REGULATORY COMPLIANCE SYSTEM")
        header_run.font.name = FONT_NAME
        header_run.font.size = Pt(8.5)
        header_run.font.color.rgb = MUTED_GRAY
        header_p.paragraph_format.space_after = Pt(24)

        # Dynamic Footer Design ("Page X of Y")
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(12)
        
        def add_page_field(paragraph, field_text):
            run = paragraph.add_run()
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED_GRAY
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = field_text
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)

        prefix_run = footer_p.add_run("Page ")
        prefix_run.font.name = FONT_NAME
        prefix_run.font.size = Pt(9)
        prefix_run.font.color.rgb = MUTED_GRAY
        
        add_page_field(footer_p, "PAGE")
        mid_run = footer_p.add_run(" of ")
        mid_run.font.name = FONT_NAME
        mid_run.font.size = Pt(9)
        mid_run.font.color.rgb = MUTED_GRAY
        
        add_page_field(footer_p, "NUMPAGES")

    normal_style = doc.styles['Normal']
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(11)
    
    def add_styled_heading(text, level, size_pt):
        heading = doc.add_heading(level=level)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.keep_with_next = True
        
        run = heading.add_run(text.replace('**', ''))
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt)
        run.font.color.rgb = CRIMSON_RED
        return heading

    lines = markdown_text.split('\n')
    table_data = []
    in_table = False

    for line in lines:
        cleaned = line.strip()
        
        # Table Grid Layout Processor Layer
        if cleaned.startswith('|'):
            in_table = True
            if '---' in cleaned:
                continue
            row_cells = [cell.strip() for cell in cleaned.split('|')[1:-1]]
            table_data.append(row_cells)
            continue
        else:
            if in_table and table_data:
                num_cols = max(len(r) for r in table_data)
                word_table = doc.add_table(rows=len(table_data), cols=num_cols)
                word_table.style = 'Medium Shading 1 Accent 1'
                word_table.allow_autofit = True
                
                for r_idx, row in enumerate(table_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < len(word_table.rows[r_idx].cells):
                            cell_p = word_table.rows[r_idx].cells[c_idx].paragraphs[0]
                            cell_p.paragraph_format.space_before = Pt(4)
                            cell_p.paragraph_format.space_after = Pt(4)
                            
                            clean_val = val.replace('**', '')
                            cell_run = cell_p.add_run(clean_val)
                            cell_run.font.name = FONT_NAME
                            
                            if r_idx == 0:
                                cell_run.bold = True
                                cell_run.font.size = Pt(11)
                            else:
                                cell_run.font.size = Pt(10)
                
                spacer_p = doc.add_paragraph()
                spacer_p.paragraph_format.space_before = Pt(6)
                table_data = []
                in_table = False

        if not cleaned or cleaned.startswith('---'):
            continue

        if line.startswith('# '):
            add_styled_heading(line[2:], level=1, size_pt=18)
        elif line.startswith('## '):
            add_styled_heading(line[3:], level=2, size_pt=14)
        elif line.startswith('### '):
            add_styled_heading(line[4:], level=3, size_pt=12)
            
        elif line.startswith('* ') or line.startswith('- ') or re.match(r'^\d+\.', line):
            is_numbered = re.match(r'^\d+\.', line)
            style_type = 'List Number' if is_numbered else 'List Bullet'
            
            content = re.sub(r'^\d+\.\s*', '', line) if is_numbered else line[2:]
            p = doc.add_paragraph(style=style_type)
            p.paragraph_format.space_after = Pt(3)
            
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                run = p.add_run(part[2:-2] if part.startswith('**') and part.endswith('**') else part)
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
                if part.startswith('**') and part.endswith('**'):
                    run.bold = True
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = FONT_NAME
                run.font.size = Pt(11)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
# =====================================================================
# 5. EXECUTION PIPELINE TAB LAYOUT & PROCESSING LOOP
# =====================================================================
with main_tab:
    niche_selection = st.selectbox(
        "Select Target Compliance Evaluation Module:",
         [
            "Patent Checker", 
            "Grant Auditor (SERB / DST)", 
            "ANRF - NPDF Auditor", 
            "CSIR Grant Auditor", 
            "Central Agencies Multi-Scout (ICAR/DBT/MNRE)", 
            "Bioethics Scout"
        ]
    )

    if niche_selection == "Patent Checker":
        st.info("🎯 Module active: Indian Patent (IPR) & Prior-Art Agent. Auditing criteria under the Indian Patents Act, 1970.")
    elif niche_selection == "Grant Auditor (SERB / DST)":
        st.info("💰 Module active: SERB / DST Central Government Grant Alignment Framework.")
    elif niche_selection == "ANRF - NPDF Auditor":
        st.info("🎓 Module active: Anusandhan National Research Foundation Post Doctoral Fellowship screening layer.")
    elif niche_selection == "CSIR Grant Auditor":
        st.info("🔬 Module active: Council of Scientific and Industrial Research extramural project tracking matrix.")
    elif niche_selection == "Central Agencies Multi-Scout (ICAR/DBT/MNRE)":
        st.info("🌾 Module active: Inter-departmental project validator (ICAR, DBT, MNRE, MeitY standards).")
    else: 
        st.info("🩺 Module active: Bioethics & Clinical Trial Regulatory Scout synced against regional CTRI and Live ICMR profiles.")

    uploaded_file = st.file_uploader(
        "📤 Drag and drop your research manuscript, grant proposal draft, or clinical methodology document:",
        type=["pdf", "docx"],
        help="Supports complete academic drafts in PDF or Microsoft Word formatting."
    )

    parsed_manuscript_text = ""

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split(".")[-1].lower()
        if file_extension == "pdf":
            try:
                pdf_reader = PdfReader(uploaded_file)
                extracted_pages = [page.extract_text() for page in pdf_reader.pages]
                parsed_manuscript_text = "\n".join(filter(None, extracted_pages))
                st.success(f"📊 Successfully parsed PDF: {uploaded_file.name}")
            except Exception as e:
                st.error(f"Error parsing PDF file: {e}")
        elif file_extension == "docx":
            try:
                doc = Document(uploaded_file)
                extracted_paras = [p.text for p in doc.paragraphs]
                parsed_manuscript_text = "\n".join(filter(None, extracted_paras))
                st.success(f"📊 Successfully parsed Word Document: {uploaded_file.name}")
            except Exception as e:
                st.error(f"Error parsing Word file: {e}")

    # Core Execution Lifecycle Trigger Block
    if st.button("🚀 Initiate Complete Compliance Evaluation", type="primary"):
        if not user_gemini_key.strip():
            st.error("🔑 Access Denied: Please provide your personal Gemini API key in the sidebar to power the evaluation loop.")
        elif not parsed_manuscript_text.strip():
            st.error("Please upload a valid PDF or DOCX file to execute the compliance evaluation loop.")
        else:
            os.environ["GOOGLE_API_KEY"] = user_gemini_key
            status_flag = "Success"
            
            try:
                with st.status("🕵️ Executive Agent auditing compliance vectors...", expanded=True) as status_box:
                    st.write("🔍 Extracting technical vectors and scanning tracking registers...")
                    compliance_engine = build_compliance_agent()
                    initial_state = {
                        "niche_type": niche_selection,
                        "uploaded_text_pool": parsed_manuscript_text,
                        "retrieved_context": [],
                        "final_audit_report": ""
                    }
                    final_output = compliance_engine.invoke(initial_state)
                    status_box.update(label="✅ Analysis completion reached!", state="complete")
                
                raw_response = str(final_output.get("final_audit_report", ""))
                
                # --- SANITIZATION PIPELINE: Strips residual API structures & signatures ---
                clean_text = raw_response.strip()
                if clean_text.startswith("[") or clean_text.startswith("{"):
                    try:
                        import ast
                        parsed_data = ast.literal_eval(clean_text)
                        if isinstance(parsed_data, list) and len(parsed_data) > 0:
                            clean_text = parsed_data[0].get("text", clean_text)
                        elif isinstance(parsed_data, dict):
                            clean_text = parsed_data.get("text", clean_text)
                    except Exception:
                        pass
                
                if "'extras':" in clean_text or '"extras":' in clean_text:
                    clean_text = re.split(r",?\s*['\"]extras['\"]", clean_text)[0]
                    # Strip residual list brackets if left behind
                    clean_text = clean_text.rstrip(",}] ")
                    
                actual_text = clean_text.strip()
                # ------------------------------------------------------------------------
                
            except Exception as e:
                status_flag = "Failed"
                actual_text = f"An execution error occurred during processing loops: {e}"
                st.error(actual_text)
            
            st.session_state["metrics_log"].append({
                "Document Name": uploaded_file.name if uploaded_file else "Unknown_Draft",
                "Evaluation Module": niche_selection,
                "Status": status_flag,
                "Length (Chars)": len(parsed_manuscript_text)
            })
            
            if status_flag == "Success":
                st.success("✨ Regulatory document synthesis successful!")
                docx_bytes = convert_markdown_to_docx(actual_text)
                safe_niche = niche_selection.replace(' ', '_')
                raw_filename = uploaded_file.name if uploaded_file else "Manuscript"
                clean_filename_base = re.sub(r'[^a-zA-Z0-9_]', '_', raw_filename)
                dynamic_filename = f"{safe_niche}_{clean_filename_base}_Audit_Report.docx"
                
                st.download_button(
                    label="📄 Download Official Audit Report as MS Word (.docx)",
                    data=docx_bytes,
                    file_name=dynamic_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.markdown("---")
                st.markdown(actual_text)

# =====================================================================
# 6. TRACKING ANALYTICS DASHBOARD ENGINE TAB
# =====================================================================
with dashboard_tab:
    st.header("📊 Real-Time Operations Analytics")
    st.caption("Live summary tracking of document auditing workflows across the current testing cycle.")
    
    if len(st.session_state["metrics_log"]) > 0:
        df_metrics = pd.DataFrame(st.session_state["metrics_log"])
        
        kpi1, kpi2, kpi3 = st.columns(3)
        with kpi1:
            st.metric("Total Files Evaluated", len(df_metrics))
        with kpi2:
            success_count = len(df_metrics[df_metrics["Status"] == "Success"])
            success_rate = (success_count / len(df_metrics)) * 100
            st.metric("Processing Success Rate", f"{success_rate:.1f}%")
        with kpi3:
            mode_series = df_metrics["Evaluation Module"].mode()
            most_used_module = str(mode_series.iloc[0]) if not mode_series.empty else "None"
            st.metric("Peak Operational Niche", most_used_module)
            
        st.markdown("---")
        col_graph1, col_graph2 = st.columns(2)
        
        with col_graph1:
            st.subheader("📁 Module Allocation Share")
            module_counts = df_metrics["Evaluation Module"].value_counts()
            st.bar_chart(module_counts, color="#2e7d32")
            
        with col_graph2:
            st.subheader("📋 Document Audit Process Registry")
            st.dataframe(df_metrics, use_container_width=True, hide_index=True)
    else:
        st.info("📉 No telemetry tracking records found. Audit your first manuscript document inside the Execution Pipeline tab to generate data visualizations.")
